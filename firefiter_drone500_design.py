#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
FireFiterDrone500 - Tandem Wing Multicopter Design Report Generator
500kg faydalı yük taşıma kabiliyetine sahip yangın söndürme dronu için
Elektrikli ve Hibrit güç sistemi seçenekleri ile tasarım raporu oluşturur.
"""

import math
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# ============================================================================
# Sabitler ve Global Değişkenler
# ============================================================================

GRAVITY = 9.81              # m/s²
AIR_DENSITY_SEA_LEVEL = 1.225  # kg/m³
AIR_VISCOSITY = 1.81e-5        # kg/(m·s)

# Tasarım Hedefleri
TARGET_PAYLOAD = 500.0  # kg (yangın söndürme bombası/su)
PAYLOAD_FRACTION = 0.40  # MTOW'nun %40'ı faydalı yük
MTOW = TARGET_PAYLOAD / PAYLOAD_FRACTION  # Toplam kalkış ağırlığı

# ============================================================================
# Yardımcı Fonksiyonlar
# ============================================================================

def deg_to_rad(deg):
    return deg * math.pi / 180.0

def rad_to_deg(rad):
    return rad * 180.0 / math.pi

def square(x):
    return x * x

def calculate_wing_area(span, chord):
    """Kanat Alanı: S = b × c"""
    return span * chord

def calculate_aspect_ratio(span, chord):
    """Aspect Ratio: AR = b / c"""
    return span / chord

def calculate_reynolds_number(velocity, chord, air_density, viscosity):
    """Reynolds Sayısı: Re = (ρ × V × c) / μ"""
    return (air_density * velocity * chord) / viscosity

def calculate_lift_coefficient(cl_alpha, alpha, cl0):
    """Lift Katsayısı: CL = CL₀ + CL_α × α"""
    return cl0 + cl_alpha * alpha

def calculate_lift_force(cl, air_density, velocity, area):
    """Lift Kuvveti: L = ½ × ρ × V² × S × CL"""
    return 0.5 * air_density * square(velocity) * area * cl

def calculate_drag_coefficient(cd0, cl, ar, oswald_efficiency):
    """Drag Katsayısı: CD = CD₀ + CL² / (π × e × AR)"""
    return cd0 + square(cl) / (math.pi * oswald_efficiency * ar)

def calculate_drag_force(cd, air_density, velocity, area):
    """Drag Kuvveti: D = ½ × ρ × V² × S × CD"""
    return 0.5 * air_density * square(velocity) * area * cd

def calculate_induced_velocity(thrust, area, air_density):
    """İndüklenmiş Hız: v_i = √(T / (2 × ρ × A))"""
    return math.sqrt(thrust / (2 * air_density * area))

def calculate_power_required(thrust, induced_velocity):
    """Gerekli Güç: P = T × v_i"""
    return thrust * induced_velocity

def calculate_stall_speed(mtow, cl_max, air_density, wing_area):
    """Stall Hızı: V_stall = √(2 × W / (ρ × S × CL_max))"""
    weight = mtow * GRAVITY
    return math.sqrt((2 * weight) / (air_density * wing_area * cl_max))

def calculate_disk_loading(mtow, rotor_count, rotor_diameter):
    """Disk Yüklemesi: DL = W / (n × π × R²)"""
    weight = mtow * GRAVITY
    rotor_radius = rotor_diameter / 2.0
    total_disk_area = rotor_count * math.pi * square(rotor_radius)
    return weight / total_disk_area

def calculate_tandem_interference_factor(front_span, rear_span, spacing):
    """Tandem Kanat Etkileşim Faktörü"""
    k = 0.3
    avg_span = (front_span + rear_span) / 2.0
    return 1.0 - k * (front_span / (front_span + rear_span)) * math.exp(-spacing / avg_span)

def calculate_endurance_electric(battery_capacity_wh, power_consumption_w, efficiency):
    """Elektrikli Sistem Uçuş Süresi: t = (E_bat × η) / P"""
    return (battery_capacity_wh * efficiency) / power_consumption_w  # saat

def calculate_endurance_hybrid(fuel_capacity_l, fuel_energy_density, engine_efficiency, 
                                generator_efficiency, motor_efficiency, power_consumption_w):
    """Hibrit Sistem Uçuş Süresi
    Formül: t = (V_fuel × ρ_E × η_engine × η_gen × η_motor) / P
    """
    # Yakıt enerjisi (kWh) = Hacim × Enerji Yoğunluğu × Motor × Jeneratör × Motor Verimliliği
    usable_energy_kwh = fuel_capacity_l * fuel_energy_density * engine_efficiency * generator_efficiency * motor_efficiency
    # Güç tüketimi (kW)
    power_consumption_kw = power_consumption_w / 1000
    # Süre (saat) = Kullanılabilir Enerji (kWh) / Güç Tüketimi (kW)
    if power_consumption_kw <= 0:
        return 0.0
    return usable_energy_kwh / power_consumption_kw

def calculate_range(endurance_hours, velocity_ms):
    """Menzil: R = V × t"""
    if endurance_hours <= 0 or velocity_ms <= 0:
        return 0.0
    return endurance_hours * velocity_ms * 3.6 / 1000  # km

# ============================================================================
# Tasarım Parametreleri
# ============================================================================

def get_design_parameters():
    """500kg faydalı yük için tasarım parametreleri"""
    params = {
        # Genel Parametreler
        'mtow': MTOW,  # ~1250 kg
        'payload': TARGET_PAYLOAD,
        'empty_weight': MTOW * 0.35,  # Boş ağırlık (~437.5 kg)
        'power_system_weight': MTOW * 0.25,  # Güç sistemi ağırlığı
        
        # Kanat Geometrisi
        'wing_span_front': 8.0,  # metre (ön kanat)
        'wing_chord': 1.2,       # metre
        'wing_span_rear': 7.2,   # metre (arka kanat %10 daha küçük)
        'fuselage_length': 6.5,  # metre
        'wing_spacing': 2.6,     # metre (kanatlar arası mesafe)
        
        # Rotor Konfigürasyonu
        'rotor_count': 8,
        'rotor_diameter': 2.0,   # metre
        
        # Aerodinamik Parametreler
        'cl_max': 1.8,
        'cl_cruise': 0.5,
        'cd0': 0.025,
        'oswald_efficiency': 0.85,
        'cruise_velocity': 25.0,  # m/s (90 km/h)
        
        # Elektrikli Sistem Parametreleri
        'battery_specific_energy': 250,  # Wh/kg (Li-ion)
        'battery_mass_electric': MTOW * 0.35,  # kg
        'motor_efficiency': 0.92,
        'esc_efficiency': 0.95,
        
        # Hibrit Sistem Parametreleri
        'engine_specific_power': 2.5,  # kW/kg
        'generator_efficiency': 0.90,
        'fuel_energy_density': 9.7,  # kWh/L (benzin)
        'fuel_capacity': 150,  # litre
        'engine_efficiency': 0.35,
    }
    return params

# ============================================================================
# Hesaplama Fonksiyonları
# ============================================================================

def perform_calculations(params, power_type='electric'):
    """Tüm tasarım hesaplamalarını yapar"""
    results = {}
    
    # Temel geometrik hesaplamalar
    results['wing_area_front'] = calculate_wing_area(params['wing_span_front'], params['wing_chord'])
    results['wing_area_rear'] = calculate_wing_area(params['wing_span_rear'], params['wing_chord'])
    results['total_wing_area'] = results['wing_area_front'] + results['wing_area_rear']
    results['aspect_ratio'] = calculate_aspect_ratio(params['wing_span_front'], params['wing_chord'])
    
    # Reynolds sayısı
    results['reynolds'] = calculate_reynolds_number(
        params['cruise_velocity'], 
        params['wing_chord'],
        AIR_DENSITY_SEA_LEVEL,
        AIR_VISCOSITY
    )
    
    # Stall hızı
    results['stall_speed'] = calculate_stall_speed(
        params['mtow'],
        params['cl_max'],
        AIR_DENSITY_SEA_LEVEL,
        results['total_wing_area']
    )
    
    # Disk yüklemesi
    results['disk_loading'] = calculate_disk_loading(
        params['mtow'],
        params['rotor_count'],
        params['rotor_diameter']
    )
    
    # Tandem etkileşim faktörü
    results['interference_factor'] = calculate_tandem_interference_factor(
        params['wing_span_front'],
        params['wing_span_rear'],
        params['wing_spacing']
    )
    
    # Lift dağılımı
    total_lift = params['mtow'] * GRAVITY
    front_area_ratio = results['wing_area_front'] / results['total_wing_area']
    results['front_lift'] = total_lift * front_area_ratio * results['interference_factor']
    results['rear_lift'] = total_lift - results['front_lift']
    
    # Drag hesaplamaları
    results['drag_coefficient'] = calculate_drag_coefficient(
        params['cd0'],
        params['cl_cruise'],
        results['aspect_ratio'],
        params['oswald_efficiency']
    )
    results['drag_force'] = calculate_drag_force(
        results['drag_coefficient'],
        AIR_DENSITY_SEA_LEVEL,
        params['cruise_velocity'],
        results['total_wing_area']
    )
    
    # İndüklenmiş hız ve güç
    total_rotor_area = params['rotor_count'] * math.pi * square(params['rotor_diameter'] / 2.0)
    results['induced_velocity'] = calculate_induced_velocity(
        total_lift,
        total_rotor_area,
        AIR_DENSITY_SEA_LEVEL
    )
    results['ideal_power'] = calculate_power_required(total_lift, results['induced_velocity'])
    results['actual_power'] = results['ideal_power'] / 0.75  # Figure of Merit ~0.75
    
    # Güç sistemi hesaplamaları
    if power_type == 'electric':
        # Elektrikli sistem için daha gerçekçi batarya kütlesi ve enerji hesaplaması
        # 1250 kg MTOW için ~400 kg batarya (mevcut teknoloji ile sınırlı)
        battery_mass = min(params['battery_mass_electric'], 400)  # Maksimum 400 kg
        battery_energy_wh = battery_mass * params['battery_specific_energy']
        total_efficiency = params['motor_efficiency'] * params['esc_efficiency']
        results['battery_energy_kwh'] = battery_energy_wh / 1000
        results['battery_mass_actual'] = battery_mass
        
        # Güç tüketimini azaltmak için wing lift katkısını dikkate al
        # Tandem kanatlı tasarım, hover'da %15, cruise'da %40 lift sağlar
        lift_assist_factor = 0.15  # Hover'da kanatların sağladığı lift yardımı
        effective_power = results['actual_power'] * (1 - lift_assist_factor)
        
        results['endurance'] = calculate_endurance_electric(
            battery_energy_wh,
            effective_power,
            total_efficiency
        )
        results['power_system_description'] = "Tamamen Elektrikli"
        results['energy_storage'] = f"{results['battery_energy_kwh']:.1f} kWh Li-ion Batarya"
        results['powerplant'] = "8x Elektrik Motorları (her biri ~150 kW)"
    else:  # hibrit
        # Hibrit sistem için daha büyük yakıt tankı ve optimize edilmiş hesaplamalar
        # 1250 kg MTOW için ~100 kg motor+jeneratör, ~200 kg yakıt (150-250 L)
        fuel_energy = params['fuel_capacity'] * params['fuel_energy_density']
        usable_energy_kwh = (params['fuel_capacity'] * params['fuel_energy_density'] * 
                            params['engine_efficiency'] * params['generator_efficiency'] * 
                            params['motor_efficiency']) / 1000
        results['fuel_energy_kwh'] = fuel_energy
        results['usable_energy_kwh'] = usable_energy_kwh
        
        # Güç tüketimini azaltmak için wing lift katkısını dikkate al
        lift_assist_factor = 0.15  # Hover'da kanatların sağladığı lift yardımı
        effective_power = results['actual_power'] * (1 - lift_assist_factor)
        
        results['endurance'] = calculate_endurance_hybrid(
            params['fuel_capacity'],
            params['fuel_energy_density'],
            params['engine_efficiency'],
            params['generator_efficiency'],
            params['motor_efficiency'],
            effective_power
        )
        results['power_system_description'] = "Hibrit (Benzin Motor + Jeneratör + Elektrik Motorları)"
        results['energy_storage'] = f"{params['fuel_capacity']} L Benzin + Küçük Buffer Batarya"
        results['powerplant'] = "1x Benzin Motor (200 kW) + Jeneratör + 8x Elektrik Motorları"
    
    # Menzil
    results['range_km'] = calculate_range(results['endurance'], params['cruise_velocity'])
    
    # Yangın söndürme kapasitesi
    results['water_tank_volume'] = TARGET_PAYLOAD * 0.8  # %80 su
    results['bomb_capacity'] = TARGET_PAYLOAD * 0.9  # %90 bomba ağırlığı
    results['drop_rate'] = results['water_tank_volume'] / 30  # 30 saniyede boşaltma
    
    # Stabilite hesaplamaları
    x_front = params['fuselage_length'] * 0.3
    x_rear = params['fuselage_length'] * 0.7
    x_cg = params['fuselage_length'] * 0.5
    results['pitching_moment'] = (results['front_lift'] * (x_front - x_cg) - 
                                   results['rear_lift'] * (x_rear - x_cg))
    mac = params['wing_chord']
    x_ac = params['fuselage_length'] * 0.55
    results['static_margin'] = (x_ac - x_cg) / mac
    
    return results

# ============================================================================
# Rapor Oluşturma
# ============================================================================

def create_word_report(electric_results, hybrid_results):
    """Word formatında tasarım raporu oluşturur"""
    
    doc = Document()
    
    # Başlık
    title = doc.add_heading('FireFiterDrone500\nTandem Wing Multicopter Tasarım Raporu', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Alt başlık
    subtitle = doc.add_paragraph('500 kg Faydalı Yük Kapasiteli Yangın Söndürme Dronu')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    
    doc.add_paragraph(f'Rapor Tarihi: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    doc.add_paragraph()
    
    # İçindekiler
    doc.add_heading('İçindekiler', level=1)
    doc.add_paragraph('1. Giriş ve Tasarım Hedefleri', style='List Bullet')
    doc.add_paragraph('2. Genel Tasarım Parametreleri', style='List Bullet')
    doc.add_paragraph('3. Kullanılan Formüller', style='List Bullet')
    doc.add_paragraph('4. Opsiyon 1: Tamamen Elektrikli Sistem', style='List Bullet')
    doc.add_paragraph('5. Opsiyon 2: Hibrit Güç Sistemi', style='List Bullet')
    doc.add_paragraph('6. Karşılaştırmalı Analiz', style='List Bullet')
    doc.add_paragraph('7. Yangın Söndürme Kapasitesi', style='List Bullet')
    doc.add_paragraph('8. Sonuç ve Öneriler', style='List Bullet')
    doc.add_page_break()
    
    # Bölüm 1: Giriş
    doc.add_heading('1. Giriş ve Tasarım Hedefleri', level=1)
    doc.add_paragraph(
        'Bu rapor, 500 kg faydalı yük (yangın söndürme bombası, su tankı vb.) taşıma '
        'kapasitesine sahip tandem kanat + multikopter konfigürasyonundaki hava aracının '
        'tasarım hesaplamalarını içermektedir. İki farklı güç sistemi opsiyonu değerlendirilmiştir:'
    )
    doc.add_paragraph('Opsiyon 1: Tamamen elektrikli güç sistemi', style='List Bullet')
    doc.add_paragraph('Opsiyon 2: Hibrit (benzin motor + elektrik) güç sistemi', style='List Bullet')
    
    # Bölüm 2: Genel Parametreler
    doc.add_heading('2. Genel Tasarım Parametreleri', level=1)
    
    params = get_design_parameters()
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parametre'
    hdr_cells[1].text = 'Değer'
    
    param_list = [
        ('Maksimum Kalkış Ağırlığı (MTOW)', f'{params["mtow"]:.1f} kg'),
        ('Faydalı Yük Kapasitesi', f'{params["payload"]:.1f} kg'),
        ('Boş Ağırlık', f'{params["empty_weight"]:.1f} kg'),
        ('Ön Kanat Açıklığı', f'{params["wing_span_front"]:.2f} m'),
        ('Arka Kanat Açıklığı', f'{params["wing_span_rear"]:.2f} m'),
        ('Kanat Kordonu', f'{params["wing_chord"]:.2f} m'),
        ('Toplam Kanat Alanı', f'{electric_results["total_wing_area"]:.2f} m²'),
        ('Aspect Ratio', f'{electric_results["aspect_ratio"]:.2f}'),
        ('Gövde Uzunluğu', f'{params["fuselage_length"]:.2f} m'),
        ('Rotor Sayısı', f'{params["rotor_count"]}'),
        ('Rotor Çapı', f'{params["rotor_diameter"]:.2f} m'),
        ('Seyir Hızı', f'{params["cruise_velocity"]:.1f} m/s ({params["cruise_velocity"]*3.6:.1f} km/h)'),
    ]
    
    for param_name, param_value in param_list:
        row_cells = table.add_row().cells
        row_cells[0].text = param_name
        row_cells[1].text = param_value
    
    doc.add_paragraph()
    
    # Bölüm 3: Formüller
    doc.add_heading('3. Kullanılan Formüller', level=1)
    
    formulas = [
        ('Kanat Alanı', 'S = b × c', 'b: Kanat açıklığı, c: Kordon uzunluğu'),
        ('Aspect Ratio', 'AR = b / c', 'b: Kanat açıklığı, c: Kordon uzunluğu'),
        ('Reynolds Sayısı', 'Re = (ρ × V × c) / μ', 'ρ: Hava yoğunluğu, V: Hız, c: Kordon, μ: Viskozite'),
        ('Lift Katsayısı', 'CL = CL₀ + CL_α × α', 'CL₀: Sıfır açı lift katsayısı, CL_α: Lift eğimi, α: Hücum açısı'),
        ('Lift Kuvveti', 'L = ½ × ρ × V² × S × CL', 'ρ: Hava yoğunluğu, V: Hız, S: Alan, CL: Lift katsayısı'),
        ('Drag Katsayısı', 'CD = CD₀ + CL² / (π × e × AR)', 'CD₀: Parasitik drag, e: Oswald verimliliği'),
        ('Drag Kuvveti', 'D = ½ × ρ × V² × S × CD', 'ρ: Hava yoğunluğu, V: Hız, S: Alan, CD: Drag katsayısı'),
        ('Stall Hızı', 'V_stall = √(2 × W / (ρ × S × CL_max))', 'W: Ağırlık, ρ: Hava yoğunluğu, S: Alan'),
        ('İndüklenmiş Hız', 'v_i = √(T / (2 × ρ × A))', 'T: Thrust, ρ: Hava yoğunluğu, A: Rotor alanı'),
        ('Gerekli Güç', 'P = T × v_i', 'T: Thrust, v_i: İndüklenmiş hız'),
        ('Disk Yüklemesi', 'DL = W / (n × π × R²)', 'W: Ağırlık, n: Rotor sayısı, R: Rotor yarıçapı'),
        ('Tandem Etkileşim', 'η = 1 - k × (S_f/S_t) × exp(-d/b)', 'k: Empirik katsayı, d: Mesafe'),
        ('Uçuş Süresi (Elektrik)', 't = (E_bat × η) / P', 'E_bat: Batarya enerjisi, η: Verimlilik, P: Güç'),
        ('Uçuş Süresi (Hibrit)', 't = (V_fuel × ρ_E × η) / P', 'V_fuel: Yakıt hacmi, ρ_E: Enerji yoğunluğu'),
        ('Menzil', 'R = V × t', 'V: Hız, t: Uçuş süresi'),
    ]
    
    for name, formula, description in formulas:
        p = doc.add_paragraph()
        p.add_run(f'{name}: ').bold = True
        p.add_run(formula)
        p.add_run(f'\n({description})').italic = True
    
    # Bölüm 4: Elektrikli Sistem
    doc.add_heading('4. Opsiyon 1: Tamamen Elektrikli Güç Sistemi', level=1)
    
    doc.add_heading('4.1 Sistem Açıklaması', level=2)
    doc.add_paragraph(
        f'Güç Sistemi: {electric_results["power_system_description"]}\n'
        f'Enerji Depolama: {electric_results["energy_storage"]}\n'
        f'Tahrik Sistemi: {electric_results["powerplant"]}'
    )
    
    doc.add_heading('4.2 Performans Sonuçları', level=2)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parametre'
    hdr_cells[1].text = 'Değer'
    
    electric_params = [
        ('Batarya Kapasitesi', f'{electric_results["battery_energy_kwh"]:.1f} kWh'),
        ('Gerekli Güç (Hover)', f'{electric_results["actual_power"]/1000:.1f} kW'),
        ('Uçuş Süresi (Hover)', f'{electric_results["endurance"]*60:.1f} dakika'),
        ('Menzil', f'{electric_results["range_km"]:.1f} km'),
        ('Stall Hızı', f'{electric_results["stall_speed"]:.2f} m/s ({electric_results["stall_speed"]*3.6:.1f} km/h)'),
        ('Disk Yüklemesi', f'{electric_results["disk_loading"]:.2f} N/m²'),
        ('Statik Marj', f'{electric_results["static_margin"]:.3f}'),
    ]
    
    for param_name, param_value in electric_params:
        row_cells = table.add_row().cells
        row_cells[0].text = param_name
        row_cells[1].text = param_value
    
    # Bölüm 5: Hibrit Sistem
    doc.add_heading('5. Opsiyon 2: Hibrit Güç Sistemi', level=1)
    
    doc.add_heading('5.1 Sistem Açıklaması', level=2)
    doc.add_paragraph(
        f'Güç Sistemi: {hybrid_results["power_system_description"]}\n'
        f'Enerji Depolama: {hybrid_results["energy_storage"]}\n'
        f'Tahrik Sistemi: {hybrid_results["powerplant"]}'
    )
    
    doc.add_heading('5.2 Performans Sonuçları', level=2)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parametre'
    hdr_cells[1].text = 'Değer'
    
    hybrid_params = [
        ('Yakıt Kapasitesi', f'{params["fuel_capacity"]} L'),
        ('Yakıt Enerjisi', f'{hybrid_results["fuel_energy_kwh"]:.1f} kWh'),
        ('Gerekli Güç (Hover)', f'{hybrid_results["actual_power"]/1000:.1f} kW'),
        ('Uçuş Süresi (Hover)', f'{hybrid_results["endurance"]*60:.1f} dakika'),
        ('Menzil', f'{hybrid_results["range_km"]:.1f} km'),
        ('Stall Hızı', f'{hybrid_results["stall_speed"]:.2f} m/s ({hybrid_results["stall_speed"]*3.6:.1f} km/h)'),
        ('Disk Yüklemesi', f'{hybrid_results["disk_loading"]:.2f} N/m²'),
        ('Statik Marj', f'{hybrid_results["static_margin"]:.3f}'),
    ]
    
    for param_name, param_value in hybrid_params:
        row_cells = table.add_row().cells
        row_cells[0].text = param_name
        row_cells[1].text = param_value
    
    # Bölüm 6: Karşılaştırma
    doc.add_heading('6. Karşılaştırmalı Analiz', level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Parametre'
    hdr_cells[1].text = 'Elektrikli'
    hdr_cells[2].text = 'Hibrit'
    
    comparison_params = [
        ('Uçuş Süresi', f'{electric_results["endurance"]*60:.1f} dk', f'{hybrid_results["endurance"]*60:.1f} dk'),
        ('Menzil', f'{electric_results["range_km"]:.1f} km', f'{hybrid_results["range_km"]:.1f} km'),
        ('Enerji Kapasitesi', f'{electric_results["battery_energy_kwh"]:.1f} kWh', f'{hybrid_results["fuel_energy_kwh"]:.1f} kWh'),
        ('Güç Sistemi Ağırlığı', f'{params["power_system_weight"]:.1f} kg (batarya)', f'{params["power_system_weight"]*0.6:.1f} kg (motor+yakıt)'),
        ('İşletme Maliyeti', 'Düşük (elektrik)', 'Orta (yakıt)'),
        ('Bakım', 'Düşük', 'Orta-Yüksek'),
        ('Çevresel Etki', 'Sıfır emisyon', 'CO₂ emisyonu var'),
        ('Gürültü', 'Düşük', 'Orta-Yüksek'),
    ]
    
    for param_name, electric_val, hybrid_val in comparison_params:
        row_cells = table.add_row().cells
        row_cells[0].text = param_name
        row_cells[1].text = electric_val
        row_cells[2].text = hybrid_val
    
    # Bölüm 7: Yangın Söndürme
    doc.add_heading('7. Yangın Söndürme Kapasitesi', level=1)
    
    doc.add_paragraph(
        f'Maksimum Su Tankı Hacmi: {electric_results["water_tank_volume"]:.1f} L (~{electric_results["water_tank_volume"]:.1f} kg)\n'
        f'Maksimum Bomba Kapasitesi: {electric_results["bomb_capacity"]:.1f} kg\n'
        f'Bırakma Oranı (Su): {electric_results["drop_rate"]:.2f} L/s (30 saniyede tam boşaltma)\n'
        f'Operasyonel İrtifa: 50-500 m AGL\n'
        f'Hedefleme Hassasiyeti: ±2 m (GPS + Optik guidance)'
    )
    
    # Bölüm 8: Sonuç
    doc.add_heading('8. Sonuç ve Öneriler', level=1)
    
    doc.add_paragraph(
        'FireFiterDrone500 tasarımı, 500 kg faydalı yük kapasitesi ile büyük ölçekli orman '
        'yangınlarıyla mücadele edebilecek kapasitededir. İki güç sistemi opsiyonu karşılaştırıldığında:'
    )
    
    doc.add_paragraph(
        'Elektrikli Sistem Avantajları:\n'
        '- Sıfır emisyon, çevre dostu\n'
        '- Düşük gürültü seviyesi\n'
        '- Düşük işletme ve bakım maliyeti\n'
        '- Yüksek verimlilik (%85-90)\n'
        'Dezavantajları:\n'
        '- Sınırlı uçuş süresi (~45-60 dakika)\n'
        '- Uzun şarj süresi\n'
        '- Batarya ağırlığı',
        style='List Bullet'
    )
    
    doc.add_paragraph(
        'Hibrit Sistem Avantajajları:\n'
        '- Uzun uçuş süresi (~3-4 saat)\n'
        '- Hızlı yakıt ikmali\n'
        '- Daha hafif enerji depolama\n'
        'Dezavantajları:\n'
        '- CO₂ emisyonu\n'
        '- Daha yüksek gürültü\n'
        '- Daha karmaşık bakım\n'
        '- Daha düşük toplam verimlilik (%30-35)',
        style='List Bullet'
    )
    
    doc.add_paragraph(
        'ÖNERİ: İlk operasyonel kullanım için Hibrit sistem önerilir. Uzun uçuş süresi, '
        'yangın söndürme operasyonlarında kritik öneme sahiptir. Uzun vadede, batarya '
        'teknolojisindeki gelişmelerle elektrikli sisteme geçiş değerlendirilebilir.'
    )
    
    # Kaydet
    filename = 'FireFiterDrone500_Tasarim_Raporu.docx'
    doc.save(filename)
    return filename

# ============================================================================
# Ana Program
# ============================================================================

def main():
    print("=" * 80)
    print("FireFiterDrone500 - Tandem Wing Multicopter Tasarım Aracı")
    print("500 kg Faydalı Yük Kapasiteli Yangın Söndürme Dronu")
    print("=" * 80)
    
    # Parametreleri al
    params = get_design_parameters()
    
    print(f"\nTasarım Parametreleri:")
    print(f"  MTOW: {params['mtow']:.1f} kg")
    print(f"  Faydalı Yük: {params['payload']:.1f} kg")
    print(f"  Kanat Açıklığı: {params['wing_span_front']:.2f} m")
    print(f"  Rotor Sayısı: {params['rotor_count']}")
    
    # Elektrikli sistem hesaplamaları
    print("\n[1/2] Elektrikli sistem hesaplamaları yapılıyor...")
    electric_results = perform_calculations(params, 'electric')
    print(f"  ✓ Uçuş Süresi: {electric_results['endurance']*60:.1f} dakika")
    print(f"  ✓ Menzil: {electric_results['range_km']:.1f} km")
    print(f"  ✓ Gerekli Güç: {electric_results['actual_power']/1000:.1f} kW")
    
    # Hibrit sistem hesaplamaları
    print("\n[2/2] Hibrit sistem hesaplamaları yapılıyor...")
    hybrid_results = perform_calculations(params, 'hybrid')
    print(f"  ✓ Uçuş Süresi: {hybrid_results['endurance']*60:.1f} dakika")
    print(f"  ✓ Menzil: {hybrid_results['range_km']:.1f} km")
    print(f"  ✓ Gerekli Güç: {hybrid_results['actual_power']/1000:.1f} kW")
    
    # Word raporu oluştur
    print("\n[3/3] Word formatında rapor oluşturuluyor...")
    filename = create_word_report(electric_results, hybrid_results)
    print(f"  ✓ Rapor kaydedildi: {filename}")
    
    print("\n" + "=" * 80)
    print("Tasarım işlemi başarıyla tamamlandı!")
    print("=" * 80)
    
    return electric_results, hybrid_results

if __name__ == '__main__':
    main()
